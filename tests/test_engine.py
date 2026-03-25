from pathlib import Path
import json
import re

from docx import Document
from PIL import Image

from agent_review.applicability import build_applicability_checks
from agent_review.adjudication import build_formal_adjudication
from agent_review.engine import TenderReviewEngine
from agent_review.extractors.clauses import extract_clauses
from agent_review.llm import QwenReviewEnhancer
from agent_review.models import (
    AdoptionStatus,
    ApplicabilityCheck,
    ApplicabilityStatus,
    ClauseRole,
    ConclusionLevel,
    Evidence,
    EvidenceBundle,
    FileType,
    FindingType,
    FormalAdjudication,
    FormalDisposition,
    Recommendation,
    ReviewMode,
    ReviewPoint,
    ReviewPointStatus,
    ReviewQualityGate,
    LegalBasis,
    QualityGateStatus,
    Severity,
)
from agent_review.outputs import write_review_artifacts
from agent_review.parsers import load_document, load_documents
from agent_review.parsers.ocr import run_ocr
from agent_review.parsers.vision_ocr import VisionOcrResult
from agent_review.rules.risk_rules import match_risk_rules
from agent_review.reporting import (
    render_formal_review_opinion,
    render_markdown,
    render_opinion_letter,
)
from agent_review.llm.prompts import build_review_point_second_review_prompt


def test_detects_manual_review_for_attachment_markers() -> None:
    text = """
    项目概况
    采购需求详见附件。
    供应商资格要求：具有相关资质。
    评分标准见附表。
    提交截止时间：2026年4月1日。
    付款方式：以正式合同为准。
    """
    report = TenderReviewEngine().review_text(text, document_name="demo.txt")

    assert any(
        item.finding_type == FindingType.manual_review_required for item in report.findings
    )
    assert report.manual_review_queue
    assert report.scope_statement
    assert report.file_info.file_type in {
        FileType.complete_tender,
        FileType.procurement_requirement,
        FileType.mixed_document,
    }
    point_map = {(item.dimension, item.title): item for item in report.review_points}
    point = point_map[("采购范围清晰度", "采购范围清晰度依赖附件或外部材料")]
    assert any(source.startswith("finding:manual_review_required:") for source in point.source_findings)
    assert point.status == ReviewPointStatus.manual_confirmation


def test_detects_restrictive_terms_warning() -> None:
    text = """
    采购需求
    本项目要求原厂服务团队，本地注册地供应商优先。
    资格要求
    供应商应具备相关资质。
    评分标准
    采用综合评分法。
    提交截止时间：2026年4月1日。
    付款方式：按合同执行。
    """
    report = TenderReviewEngine().review_text(text, document_name="demo.txt")

    assert any(item.title == "发现潜在限制性竞争表述" for item in report.findings)
    assert any(item.rule_name == "指定品牌/原厂限制" for item in report.risk_hits)
    assert any(item.legal_basis for item in report.risk_hits if item.rule_name == "指定品牌/原厂限制")
    assert not any(item.rule_name == "主观评分表述" for item in report.risk_hits)
    assert report.overall_conclusion in {ConclusionLevel.revise, ConclusionLevel.reject}


def test_detects_sme_personnel_and_contract_risks() -> None:
    text = """
    项目名称：某服务项目
    项目属性：服务
    中小企业声明函：制造商声明
    本项目专门面向中小企业采购，仍适用价格扣除。
    年龄要求35岁以下，限女性，身高160以上。
    人员更换须经采购人同意，采购人有权直接指挥现场人员。
    采购人拥有最终解释权。
    尾款根据采购人满意度考核后支付。
    """
    report = TenderReviewEngine().review_text(text, document_name="demo.txt")
    titles = {item.title for item in report.findings}

    assert "专门面向中小企业却仍保留价格扣除" in titles
    assert "服务项目声明函类型疑似错用货物模板" in titles
    assert "性别限制" in titles
    assert "采购人直接指挥" in titles
    assert "采购人单方解释或决定条款" in titles
    assert "尾款支付与考核条款联动风险" in titles


def test_detects_project_structure_and_template_conflicts() -> None:
    text = """
    项目名称：某物业服务项目
    采购标的：物业管理服务
    品目名称：办公家具
    项目属性：服务
    所属行业：工业
    中小企业声明函：制造商声明
    本项目专门面向中小企业采购，仍适用价格扣除。
    合同条款中写明质保期2年。
    """
    report = TenderReviewEngine().review_text(text, document_name="demo.txt")
    titles = {item.title for item in report.findings}

    assert "项目属性与所属行业口径疑似不一致" in titles
    assert "项目属性与声明函模板口径冲突" in titles
    assert "服务项目保留货物类声明函模板" in titles
    assert "专门面向中小企业却保留价格扣除模板" in titles
    assert "项目属性 vs 品目名称" in titles
    assert "项目属性 vs 所属行业" in titles
    assert "项目属性 vs 中小企业声明函" in titles
    assert "baseline_risk_rules" in report.rule_selection.core_modules
    assert "sme_policy" in report.rule_selection.core_modules
    assert "template_conflicts" in report.rule_selection.core_modules
    assert "project_structure" in report.rule_selection.enhancement_modules
    assert "service" in report.rule_selection.scenario_tags


def test_missing_dimension_generates_missing_evidence() -> None:
    text = "这是一份极短的文本，只提到项目概况。"
    report = TenderReviewEngine().review_text(text, document_name="short.txt")

    assert any(item.finding_type == FindingType.missing_evidence for item in report.findings)
    assert report.section_index
    assert isinstance(report.relative_strengths, list)
    assert [item.stage_name for item in report.stage_records] == [
        "document_structure",
        "clause_extraction",
        "clause_role_classification",
        "review_task_planning",
        "dimension_review",
        "rule_evaluation",
        "consistency_review",
        "review_point_assembly",
        "applicability_check",
        "review_quality_gate",
        "formal_adjudication",
        "finalize_report",
    ]
    assert all(isinstance(item.clause_role, ClauseRole) for item in report.extracted_clauses)
    assert report.review_points
    assert report.review_point_catalog
    assert report.applicability_checks
    assert report.quality_gates
    assert report.formal_adjudication


def test_review_point_catalog_covers_structured_policy_and_contract_points() -> None:
    text = """
    项目属性：服务
    中小企业声明函：制造商声明
    本项目专门面向中小企业采购，仍适用价格扣除。
    付款方式：尾款支付以考核结果为准。
    """
    report = TenderReviewEngine().review_text(text, document_name="demo.txt")

    catalog_map = {item.catalog_id: item.title for item in report.review_point_catalog}
    assert "RP-SME-001" in catalog_map
    assert "RP-SME-002" in catalog_map
    assert any(
        item.catalog_id == "RP-CONTRACT-005"
        for item in report.review_points
        if item.title == "尾款支付与考核条款联动风险"
    )


def test_review_task_planning_builds_standard_tasks_without_polluting_findings() -> None:
    text = """
    项目属性：服务
    采购标的：物业服务
    """
    report = TenderReviewEngine().review_text(text, document_name="demo.txt")

    assert any(item.stage_name == "review_task_planning" for item in report.stage_records)
    assert any(item.catalog_id == "RP-SME-002" for item in report.review_points)
    assert any(item.catalog_id == "RP-PER-001" for item in report.review_points)
    assert not any(item.title == "服务项目声明函类型疑似错用货物模板" for item in report.findings)


def test_review_task_fact_collectors_attach_structured_facts_to_tasks() -> None:
    text = """
    项目属性：服务
    中小企业声明函（货物）：全部货物由中小企业制造。
    本项目专门面向中小企业采购，仍适用价格扣除。
    付款方式：尾款于验收合格后支付，且与满意度考核结果挂钩。
    """
    report = TenderReviewEngine().review_text(text, document_name="demo.txt")

    task_map = {item.catalog_id: item for item in report.review_points}
    sme_task = task_map["RP-SME-002"]
    contract_task = task_map["RP-CONTRACT-005"]

    assert sme_task.status in {
        ReviewPointStatus.identified,
        ReviewPointStatus.suspected,
        ReviewPointStatus.manual_confirmation,
    }
    assert sme_task.evidence_bundle.direct_evidence
    assert any("项目属性=服务" in item.quote for item in sme_task.evidence_bundle.direct_evidence)
    assert any(
        "中小企业声明函类型" in item.quote
        for item in (
            sme_task.evidence_bundle.direct_evidence
            + sme_task.evidence_bundle.supporting_evidence
            + sme_task.evidence_bundle.conflicting_evidence
        )
    )

    assert contract_task.evidence_bundle.direct_evidence
    assert any("付款节点=存在" in item.quote for item in contract_task.evidence_bundle.direct_evidence)


def test_standard_task_library_and_task_specific_evidence_can_cover_scoring_and_template_tasks() -> None:
    text = """
    项目属性：货物
    采购标的：教师公寓家具
    样品要求：投标人须提供样品。
    样品分：10分
    财务指标：营业收入越高得分越高。
    不接受联合体投标，不允许合同分包。
    联合体共同投标协议书
    """
    report = TenderReviewEngine().review_text(text, document_name="demo.txt")
    task_map = {item.catalog_id: item for item in report.review_points}

    assert "RP-SCORE-003" in task_map
    assert "RP-SCORE-004" in task_map
    assert "RP-TPL-006" in task_map
    assert any("样品分" in item.quote for item in task_map["RP-SCORE-003"].evidence_bundle.direct_evidence)
    assert any("营业收入" in item.quote or "财务指标" in item.quote for item in task_map["RP-SCORE-004"].evidence_bundle.direct_evidence)
    assert task_map["RP-TPL-006"].evidence_bundle.supporting_evidence or task_map["RP-TPL-006"].evidence_bundle.conflicting_evidence


def test_task_evidence_assembler_can_capture_conflicting_and_rebuttal_evidence() -> None:
    text = """
    项目属性：服务
    法定代表人证明书
    附：代表人性别：_____年龄：_________
    本项目专门面向中小企业采购。
    价格扣除不适用本项目。
    """
    report = TenderReviewEngine().review_text(text, document_name="demo.txt")
    task_map = {item.catalog_id: item for item in report.review_points}

    personnel_task = task_map["RP-PER-001"]
    policy_task = task_map["RP-SME-001"]

    assert not personnel_task.evidence_bundle.direct_evidence
    assert personnel_task.evidence_bundle.supporting_evidence
    assert policy_task.evidence_bundle.conflicting_evidence or policy_task.evidence_bundle.rebuttal_evidence
    assert any(
        "价格扣除不适用" in note or "反证" in note or "冲突" in note
        for note in policy_task.evidence_bundle.missing_evidence_notes
    )


def test_false_positive_extractors_do_not_misclassify_template_or_ip_lines() -> None:
    text = """
    附：代表人性别：_____年龄：_________ 身份证号码：__________________
    使用过程中不会产生因第三方提出侵犯其专利权、商标权或其它知识产权而引起的纠纷。
    5.以联合体形式投标的，应符合以下规定：
    """
    report = TenderReviewEngine().review_text(text, document_name="demo.txt")
    fields = {item.field_name: item.content for item in report.extracted_clauses}
    assert "年龄限制" not in fields
    assert "容貌体形要求" not in fields
    assert "是否有限制产地厂家商标" not in fields
    assert "是否要求专利" not in fields


def test_false_positive_rules_do_not_promote_general_management_or_template_lines() -> None:
    text = """
    项目属性：服务
    附：代表人性别：_____年龄：_________ 身份证号码：__________________
    员工达到退休年龄的须购买意外险等必要的商业保险。
    设施设备维修计划报采购人审批。
    中标人在日常管理中建立量化考核并报采购人审核备案。
    """
    report = TenderReviewEngine().review_text(text, document_name="demo.txt")
    titles = {item.title for item in report.findings}
    assert "年龄限制" not in titles
    assert "采购人审批录用" not in titles
    assert "考核条款可能控制付款或履约评价" not in titles
    assert "扣款机制可能过度依赖单方考核" not in titles


def test_applicability_prefers_structured_clause_fields() -> None:
    text = """
    项目属性：服务
    中小企业声明函：制造商声明
    本项目专门面向中小企业采购，仍适用价格扣除。
    """
    report = TenderReviewEngine().review_text(text, document_name="demo.txt")

    applicability_map = {item.catalog_id: item for item in report.applicability_checks}
    sme_check = applicability_map["RP-SME-001"]
    service_template_check = applicability_map["RP-SME-002"]

    assert sme_check.applicable is True
    assert "结构化字段" in sme_check.requirement_results[0].detail
    assert sme_check.requirement_chain_complete is True
    assert "项目专门面向中小企业" in sme_check.satisfied_conditions
    assert service_template_check.applicable is True
    assert any("中小企业声明函类型" in item.detail for item in service_template_check.requirement_results)
    assert service_template_check.requirement_chain_complete is True


def test_extract_clauses_normalizes_structured_values() -> None:
    text = """
    项目属性：服务
    本项目专门面向中小企业采购。
    中小企业声明函（货物）：全部货物由中小企业制造。
    本项目仍适用价格扣除。
    本采购包不接受联合体投标，不允许合同分包。
    付款方式：尾款于验收合格后支付，且与满意度考核结果挂钩。
    """

    clauses = {item.field_name: item for item in extract_clauses(text)}

    assert clauses["项目属性"].normalized_value == "服务"
    assert clauses["是否专门面向中小企业"].normalized_value == "是"
    assert "制造商" in clauses["中小企业声明函类型"].normalized_value
    assert clauses["是否仍保留价格扣除条款"].normalized_value == "是"
    assert clauses["是否允许联合体"].normalized_value == "不允许"
    assert clauses["是否允许分包"].normalized_value == "不允许"
    assert "尾款" in clauses["付款节点"].relation_tags
    assert "考核联动" in clauses["付款节点"].relation_tags


def test_applicability_uses_structured_field_relations() -> None:
    text = """
    项目属性：服务
    中小企业声明函（货物）：全部货物由中小企业制造。
    本项目专门面向中小企业采购，仍适用价格扣除。
    付款方式：尾款于验收合格后支付。
    尾款根据满意度考核结果支付。
    """

    clauses = extract_clauses(text)
    points = [
        ReviewPoint(
            point_id="RP-T-001",
            catalog_id="RP-SME-002",
            title="服务项目声明函类型疑似错用货物模板",
            dimension="中小企业政策风险",
            severity=Severity.high,
            status=ReviewPointStatus.confirmed,
            rationale="结构化字段已抽到服务项目和制造商口径。",
            evidence_bundle=EvidenceBundle(),
        ),
        ReviewPoint(
            point_id="RP-T-002",
            catalog_id="RP-CONTRACT-005",
            title="尾款支付与考核条款联动风险",
            dimension="合同与履约风险",
            severity=Severity.high,
            status=ReviewPointStatus.confirmed,
            rationale="结构化字段已抽到尾款支付与考核联动。",
            evidence_bundle=EvidenceBundle(),
        ),
    ]
    applicability_map = {item.catalog_id: item for item in build_applicability_checks(points, clauses)}

    assert applicability_map["RP-SME-002"].applicable is True
    assert any("项目属性=服务" in item.detail for item in applicability_map["RP-SME-002"].requirement_results)
    assert applicability_map["RP-CONTRACT-005"].applicable is True
    assert any("尾款/考核联动" in item.detail for item in applicability_map["RP-CONTRACT-005"].requirement_results)


def test_applicability_requirement_chain_reports_missing_and_blocking_conditions() -> None:
    text = """
    项目属性：服务
    本项目专门面向中小企业采购。
    价格扣除不适用本项目。
    """
    report = TenderReviewEngine().review_text(text, document_name="demo.txt")
    applicability_map = {item.catalog_id: item for item in report.applicability_checks}
    policy_check = applicability_map["RP-SME-001"]

    assert policy_check.applicable is False
    assert policy_check.requirement_chain_complete is False
    assert "文件仍保留价格扣除" in policy_check.missing_conditions
    assert "存在冲突证据" in policy_check.blocking_conditions
    assert "要件链被阻断" in policy_check.summary or "要件链未闭合" in policy_check.summary


def test_forestry_like_project_can_trigger_structure_scoring_contract_and_amount_review_points() -> None:
    text = """
    项目名称：丹巴县2024年造林绿化项目
    项目属性：货物
    采购标的：苗木、肥料、防治药剂、标识牌及人工管护服务
    合同类型：承揽合同
    合同履行期限：1095日
    采购内容：清林整地、栽植、连续三年施肥、幼林抚育、成林管护、机械运水
    预算金额：2,899,600.00元
    最高限价：2,680,443.18元
    面向中小企业采购金额：2,680,443.18元
    评分标准：
    软件企业认定证书5分
    ITSS运行维护服务证书2分
    利润率10分
    财务报告2分
    实施方案30分，齐全且无缺陷得满分，每缺项扣分，每处缺陷扣2.5分
    合同条款：不得将本项目成果移作他用，不得向第三方泄露本项目成果。
    验收条款：如采购文件与投标文件约定标准抵触，由采购人按质量要求和技术指标、行业标准比较优胜的原则确定验收标准。
    """
    report = TenderReviewEngine().review_text(text, document_name="forest.txt")
    titles = {item.title for item in report.findings}
    clause_map = {}
    for item in report.extracted_clauses:
        clause_map.setdefault(item.field_name, []).append(item)
    applicability_map = {item.catalog_id: item for item in report.applicability_checks}

    assert clause_map["合同类型"][0].normalized_value == "承揽合同"
    assert clause_map["是否含持续性服务"][0].normalized_value == "是"
    assert "人工管护" in clause_map["采购内容构成"][0].relation_tags
    assert "软件企业认定证书" in clause_map["行业相关性存疑评分项"][0].normalized_value
    assert clause_map["方案评分扣分模式"][0].normalized_value == "存在"
    assert clause_map["合同成果模板术语"][0].normalized_value == "存在"
    assert clause_map["验收弹性条款"][0].normalized_value == "存在"
    assert clause_map["面向中小企业采购金额"][0].normalized_value == "2680443.18"

    assert "货物采购混入持续性作业服务" in titles
    assert "项目属性与合同类型口径疑似不一致" in titles
    assert "行业无关证书或财务指标被纳入评分" in titles
    assert "方案评分量化不足" in titles
    assert "合同条款出现非本行业成果模板表述" in titles
    assert "验收标准存在优胜原则或单方弹性判断" in titles

    assert applicability_map["RP-STRUCT-007"].applicable is True
    assert applicability_map["RP-STRUCT-008"].applicable is True
    assert applicability_map["RP-SCORE-005"].applicable is True
    assert applicability_map["RP-SCORE-006"].applicable is True
    assert applicability_map["RP-CONTRACT-008"].applicable is True
    assert applicability_map["RP-CONTRACT-009"].applicable is True
    assert applicability_map["RP-CONS-009"].applicable is True
    assert applicability_map["RP-SME-005"].applicable is True


def test_contract_type_extractor_prefers_explicit_type_over_government_procurement_contract_phrase() -> None:
    text = """
    （十三）是否属于签订不超过3年履行期限政府采购合同的项目：否
    1）合同类型：承揽合同
    """
    clauses = {item.field_name: item for item in extract_clauses(text)}

    assert clauses["合同类型"].normalized_value == "承揽合同"


def test_load_document_supports_docx(tmp_path: Path) -> None:
    file_path = tmp_path / "sample.docx"
    document = Document()
    document.add_paragraph("项目概况")
    document.add_paragraph("采购需求：提供运维服务。")
    table = document.add_table(rows=1, cols=2)
    table.rows[0].cells[0].text = "预算金额"
    table.rows[0].cells[1].text = "100000"
    document.save(file_path)

    document_name, parse_result = load_document(file_path)
    assert document_name == "sample.docx"
    assert parse_result.source_format == "docx"
    assert "项目概况" in parse_result.text
    assert parse_result.tables


def test_load_document_supports_image_ocr_with_warning_when_tesseract_missing(monkeypatch, tmp_path: Path) -> None:
    image_path = tmp_path / "sample.png"
    image = Image.new("RGB", (120, 40), color="white")
    image.save(image_path)
    monkeypatch.setattr(
        "agent_review.parsers.ocr.run_vision_ocr",
        lambda **kwargs: VisionOcrResult(
            doc_type="image",
            summary="",
            extracted_text="",
            fields={},
            confidence=None,
            warnings=["视觉 OCR 未生效: mocked"],
        ),
    )

    document_name, parse_result = load_document(image_path)
    assert document_name == "sample.png"
    assert parse_result.source_format == "png"
    assert parse_result.page_count == 1
    assert isinstance(parse_result.warnings, list)


def test_run_ocr_can_extract_table_like_rows(monkeypatch, tmp_path: Path) -> None:
    image_path = tmp_path / "table.png"
    image = Image.new("RGB", (200, 120), color="white")
    image.save(image_path)

    def fake_image_to_string(_image, lang=None):
        return "预算金额 100000\n最高限价 90000"

    def fake_image_to_data(_image, lang=None, output_type=None):
        return {
            "text": ["预算金额", "100000", "最高限价", "90000", "备注"],
            "conf": ["90", "88", "92", "87", "86"],
            "block_num": [1, 1, 1, 1, 2],
            "par_num": [1, 1, 1, 1, 1],
            "line_num": [1, 1, 2, 2, 1],
            "left": [10, 120, 10, 120, 10],
        }

    monkeypatch.setattr("agent_review.parsers.ocr.pytesseract.image_to_string", fake_image_to_string)
    monkeypatch.setattr("agent_review.parsers.ocr.pytesseract.image_to_data", fake_image_to_data)
    monkeypatch.setattr(
        "agent_review.parsers.ocr.run_vision_ocr",
        lambda **kwargs: VisionOcrResult(
            doc_type="报价表",
            summary="图片为报价表截图",
            extracted_text="",
            fields={"table_headers": ["字段", "数值"]},
            confidence=0.9,
            warnings=[],
        ),
    )

    result = run_ocr(image_path)

    assert "预算金额" in result.text
    assert result.tables
    assert result.tables[0].source == "ocr_table"
    assert result.tables[0].row_count == 2
    assert "视觉OCR摘要" in result.text


def test_load_documents_can_merge_multiple_sources(tmp_path: Path) -> None:
    first = tmp_path / "a.txt"
    second = tmp_path / "b.txt"
    first.write_text("项目属性：服务\n采购需求：物业服务。", encoding="utf-8")
    second.write_text("合同条款\n付款方式：按月支付。", encoding="utf-8")

    document_name, parse_result, source_documents = load_documents([first, second])

    assert "等2个文件" in document_name
    assert parse_result.source_format == "multi"
    assert len(source_documents) == 2
    assert "## 文档：a.txt" in parse_result.text
    assert "## 文档：b.txt" in parse_result.text


class FakeEnhancer:
    def enhance(self, report):
        report.summary = "这是经过LLM增强的结论摘要。"
        report.llm_enhanced = True
        report.specialist_tables.summaries["sme_policy"] = "这是经过LLM增强的中小企业政策专项摘要。"
        report.recommendations = [
            Recommendation(related_issue="测试问题", suggestion="这是经过LLM增强的建议。")
        ]
        return report


class FakeClient:
    def generate_text(self, system_prompt: str, user_prompt: str) -> str:
        if "补充可能遗漏但有文本依据的条款事实" in system_prompt:
            return json.dumps(
                {
                    "clause_supplements": [
                        {
                            "category": "政策条款",
                            "field_name": "分包比例",
                            "content": "文件疑似提及分包落实中小企业政策，但比例未在现有抽取结果中单列。",
                            "source_anchor": "line:8",
                            "adoption_status": "需人工确认",
                            "review_note": "分包比例出现在模糊表述中，需结合原表格确认。",
                        }
                    ]
                },
                ensure_ascii=False,
            )
        if "识别当前项目的采购场景，并补充建议的动态审查任务" in system_prompt:
            return json.dumps(
                {
                    "scenario_review_summary": "该项目呈现货物与持续性服务混合特征，建议补充结构错配类审查任务。",
                    "dynamic_review_tasks": [
                        {
                            "catalog_id": "RP-DYN-001",
                            "title": "项目属性与采购内容结构错配",
                            "dimension": "项目结构风险",
                            "severity": "high",
                            "task_type": "structure",
                            "scenario_tags": ["dynamic", "hybrid"],
                            "focus_fields": ["项目属性", "采购标的", "合同履行期限"],
                            "signal_groups": [["人工管护", "抚育", "管护"], ["承揽合同"]],
                            "evidence_hints": [
                                "优先采集项目属性、采购标的、合同履行期限和合同类型条款"
                            ],
                            "rebuttal_templates": [["仅供货", "不含人工服务"]],
                            "enhancement_fields": ["项目属性", "采购标的", "合同履行期限"],
                            "basis_hint": "当货物采购混入持续性服务和承揽口径时，应核查项目属性和合同类型是否错配。",
                        }
                    ],
                },
                ensure_ascii=False,
            )
        if "专门分析评分章节或评分相关条款的语义风险" in system_prompt:
            return json.dumps(
                {
                    "scoring_review_summary": "评分章节同时存在主观分档和证书权重偏重的风险特征，建议拆成两类评分动态任务。",
                    "dynamic_review_tasks": [
                        {
                            "catalog_id": "RP-DYN-SCORE-001",
                            "title": "评分分档主观性与量化充分性复核",
                            "dimension": "评审标准明确性",
                            "severity": "high",
                            "task_type": "scoring",
                            "scenario_tags": ["dynamic", "scoring"],
                            "focus_fields": ["评分方法", "方案评分扣分模式"],
                            "signal_groups": [["完全满足", "不完全满足"], ["证书", "财务"]],
                            "evidence_hints": [
                                "优先采集评分方法、方案评分扣分模式、行业相关性存疑评分项和采购标的"
                            ],
                            "rebuttal_templates": [["法定强制认证", "中标后提交"]],
                            "enhancement_fields": ["评分方法", "方案评分扣分模式", "行业相关性存疑评分项"],
                            "basis_hint": "评分分档主观性和证书权重偏高容易影响评审客观性。",
                        },
                        {
                            "catalog_id": "RP-DYN-SCORE-002",
                            "title": "证书检测报告及财务指标权重合理性复核",
                            "dimension": "评审标准明确性",
                            "severity": "high",
                            "task_type": "scoring",
                            "scenario_tags": ["dynamic", "scoring"],
                            "focus_fields": ["评分方法", "行业相关性存疑评分项"],
                            "signal_groups": [["证书", "检测报告"], ["财务", "分值"]],
                            "evidence_hints": [
                                "优先采集证书类评分项、检测报告要求、财务指标评分项、分值和采购标的"
                            ],
                            "rebuttal_templates": [["法定强制认证", "中标后提交"]],
                            "enhancement_fields": ["评分方法", "行业相关性存疑评分项", "财务指标加分"],
                            "basis_hint": "证书、检测报告和财务指标分值偏重时，容易加重投标负担并影响相关性。",
                        }
                    ],
                },
                ensure_ascii=False,
            )
        if "条款角色判断做复核" in system_prompt:
            return json.dumps(
                {
                    "role_review_notes": [
                        "部分审查点证据接近模板或定义说明边界，formal 前仍应结合条款角色过滤。"
                    ]
                },
                ensure_ascii=False,
            )
        if "证据包做复核" in system_prompt:
            return json.dumps(
                {
                    "evidence_review_notes": [
                        "中小企业政策冲突审查点证据较强，但付款与满意度联动类问题仍需补充合同原文。"
                    ]
                },
                ensure_ascii=False,
            )
        if "适法性判断做复核" in system_prompt:
            return json.dumps(
                {
                    "applicability_review_notes": [
                        "专门面向中小企业且保留价格扣除的要件基本满足，可直接进入 formal 审查意见。"
                    ]
                },
                ensure_ascii=False,
            )
        if "以 ReviewPoint 为单位进行二审" in system_prompt:
            return json.dumps(
                {
                    "review_point_second_reviews": [
                        {
                            "point_id": "RP-001",
                            "title": "专门面向中小企业却仍保留价格扣除",
                            "role_judgment": "当前证据来源主要是采购约束条款，角色判断基本可靠。",
                            "evidence_judgment": "已同时看到专门面向中小企业和价格扣除口径，直接证据较强。",
                            "applicability_judgment": "要件链基本闭合，可支撑政策冲突判断。",
                            "suggested_disposition": "include",
                            "rationale": "该审查点证据、角色和适法性相互支撑，可作为 formal 高风险问题保留。",
                            "adoption_status": "可直接采用",
                        }
                    ]
                },
                ensure_ascii=False,
            )
        if "补充近似但未命中的专项风险" in system_prompt:
            return json.dumps(
                {
                    "specialist_findings": [
                        {
                            "dimension": "专项语义复核",
                            "title": "评分因素与履约考核存在隐性耦合",
                            "severity": "high",
                            "rationale": "评分承诺与后续考核扣款口径疑似共用同一表述，可能导致重复约束。",
                            "source_anchor": "line:12",
                            "next_action": "拆分投标评分承诺与履约考核口径。",
                            "confidence": 0.88,
                            "adoption_status": "可直接采用",
                        }
                    ],
                    "specialist_summaries": {
                        "sme_policy": "中小企业政策专项仍存在模板与执行口径混杂问题。"
                    },
                    "recommendations": [
                        {
                            "related_issue": "评分因素与履约考核存在隐性耦合",
                            "suggestion": "拆分评审承诺与履约考核条款，避免形成双重约束。",
                        }
                    ],
                },
                ensure_ascii=False,
            )
        if "补充跨章节、跨表格、跨措辞的深层冲突" in system_prompt:
            return json.dumps(
                {
                    "consistency_findings": [
                        {
                            "dimension": "深层一致性复核",
                            "title": "付款条件与满意度表述存在隐性冲突",
                            "severity": "medium",
                            "rationale": "付款条款虽未直接写考核，但满意度表述可能实际控制尾款支付。",
                            "source_anchor": "line:15",
                            "next_action": "将付款条件改为客观验收节点。",
                            "confidence": 0.66,
                            "adoption_status": "需人工确认",
                            "review_note": "需结合合同完整上下文确认付款触发机制。",
                        }
                    ]
                },
                ensure_ascii=False,
            )
        return json.dumps(
            {
                "summary": "这是经过LLM语义复核增强后的总体结论摘要。",
                "verdict_review": "基于现有事实，除已命中规则外，仍存在评分、考核、付款联动的隐性实质性风险，建议人工重点复核。",
            },
            ensure_ascii=False,
        )


class FakeSecondReviewOverrideClient:
    def generate_text(self, system_prompt: str, user_prompt: str) -> str:
        if "以 ReviewPoint 为单位进行二审" in system_prompt:
            point_match = re.search(r'"point_id"\s*:\s*"([^"]+)"', user_prompt)
            title_match = re.search(r'"title"\s*:\s*"([^"]+)"', user_prompt)
            match = point_match.group(1) if point_match else "RP-001"
            title = title_match.group(1) if title_match else "审查点"
            return json.dumps(
                {
                    "review_point_second_reviews": [
                        {
                            "point_id": match,
                            "title": title,
                            "role_judgment": "当前证据来源需要更保守处理。",
                            "evidence_judgment": "现有证据虽可疑，但不足以直接进入正式高风险。",
                            "applicability_judgment": "要件链仍需人工补强确认。",
                            "suggested_disposition": "manual_confirmation",
                            "rationale": "LLM二审认为当前 formal 定性偏重，建议降级为待人工确认。",
                            "adoption_status": "可直接采用",
                        }
                    ]
                },
                ensure_ascii=False,
            )
        if "总体结论形成前" in system_prompt:
            return json.dumps(
                {
                    "summary": "这是经过LLM语义复核增强后的总体结论摘要。",
                    "verdict_review": "LLM二审建议将当前 formal 结论作保守处理。",
                },
                ensure_ascii=False,
            )
        return json.dumps({}, ensure_ascii=False)


class FakeDynamicTypedSecondReviewClient:
    def generate_text(self, system_prompt: str, user_prompt: str) -> str:
        if "以 ReviewPoint 为单位进行二审" in system_prompt:
            if '"task_type": "structure"' in user_prompt and "项目属性与采购内容结构错配" in user_prompt:
                return json.dumps(
                    {
                        "review_point_second_reviews": [
                            {
                                "point_id": "DYN-001",
                                "title": "项目属性与采购内容结构错配",
                                "role_judgment": "结构类动态任务当前主要来自采购约束条款，角色判断基本可靠。",
                                "evidence_judgment": "项目属性、采购标的和履约周期已进入证据链，但需结合反证审慎判断。",
                                "applicability_judgment": "结构类错配要件已有初步支撑，但仍需关注仅供货类反证。",
                                "suggested_disposition": "manual_confirmation",
                                "rationale": "结构类二审重点已触发，建议保守处理为待人工确认。",
                                "adoption_status": "可直接采用",
                            }
                        ]
                    },
                    ensure_ascii=False,
                )
            return json.dumps({"review_point_second_reviews": []}, ensure_ascii=False)
        if "总体结论形成前" in system_prompt:
            return json.dumps({"summary": "s", "verdict_review": "v"}, ensure_ascii=False)
        return json.dumps({}, ensure_ascii=False)


def test_engine_can_apply_llm_enhancer() -> None:
    text = """
    项目概况
    采购需求详见附件。
    评分标准见附表。
    """
    engine = TenderReviewEngine(review_enhancer=FakeEnhancer(), review_mode=ReviewMode.enhanced)
    report = engine.review_text(text, document_name="demo.txt")

    assert report.llm_enhanced is True
    assert report.summary == "这是经过LLM增强的结论摘要。"
    assert report.recommendations[0].suggestion == "这是经过LLM增强的建议。"
    assert report.specialist_tables.summaries["sme_policy"] == "这是经过LLM增强的中小企业政策专项摘要。"


def test_qwen_enhancer_can_merge_semantic_review_outputs() -> None:
    text = """
    项目属性：服务
    中小企业声明函：制造商声明
    评分方法：综合评分法
    方案评分扣分模式：完全满足/不完全满足。
    本项目专门面向中小企业采购，仍适用价格扣除。
    供应商需承诺服务满意度，尾款支付与履约评价挂钩。
    """
    base_report = TenderReviewEngine(review_mode=ReviewMode.fast).review_text(
        text, document_name="demo.txt"
    )
    enhancer = QwenReviewEnhancer(client=FakeClient())
    enhanced_report = enhancer.enhance(base_report)

    assert enhanced_report.llm_enhanced is True
    assert enhanced_report.llm_semantic_review.scenario_review_summary
    assert enhanced_report.llm_semantic_review.scoring_review_summary
    assert enhanced_report.llm_semantic_review.dynamic_review_tasks
    assert enhanced_report.llm_semantic_review.scoring_dynamic_review_tasks
    assert len(enhanced_report.llm_semantic_review.scoring_dynamic_review_tasks) == 2
    assert {
        item.title for item in enhanced_report.llm_semantic_review.scoring_dynamic_review_tasks
    } == {
        "评分分档主观性与量化充分性复核",
        "证书检测报告及财务指标权重合理性复核",
    }
    assert enhanced_report.llm_semantic_review.dynamic_review_tasks[0].evidence_hints
    assert enhanced_report.llm_semantic_review.dynamic_review_tasks[0].rebuttal_templates
    assert enhanced_report.llm_semantic_review.dynamic_review_tasks[0].enhancement_fields
    assert any(item.title == "项目属性与采购内容结构错配" for item in enhanced_report.review_points)
    assert enhanced_report.llm_semantic_review.review_point_second_reviews
    assert any(item.stage_name == "llm_semantic_review" for item in enhanced_report.stage_records)
    llm_tasks = {item.task_name: item.status.value for item in enhanced_report.task_records if item.task_name.startswith("llm_")}
    assert llm_tasks == {
        "llm_scenario_review": "completed",
        "llm_scoring_review": "completed",
        "llm_clause_supplement": "skipped",
        "llm_role_review": "skipped",
        "llm_evidence_review": "skipped",
        "llm_applicability_review": "skipped",
        "llm_review_point_second_review": "completed",
        "llm_specialist_review": "skipped",
        "llm_consistency_review": "skipped",
        "llm_verdict_review": "skipped",
    }
    markdown = render_markdown(enhanced_report)
    assert "## LLM场景识别与动态任务" in markdown
    assert "## LLM评分语义分析与动态任务" in markdown
    assert "证据提示" in markdown
    assert "反证模板" in markdown
    assert "## LLM审查点二审" in markdown


def test_llm_second_review_can_downgrade_formal_adjudication() -> None:
    text = """
    项目属性：服务
    本项目专门面向中小企业采购，仍适用价格扣除。
    """
    base_report = TenderReviewEngine(review_mode=ReviewMode.fast).review_text(
        text, document_name="demo.txt"
    )
    assert any(item.included_in_formal for item in base_report.formal_adjudication)

    enhancer = QwenReviewEnhancer(client=FakeSecondReviewOverrideClient())
    enhanced_report = enhancer.enhance(base_report)

    assert enhanced_report.llm_semantic_review.review_point_second_reviews
    target = next(
        item for item in enhanced_report.formal_adjudication if item.catalog_id == "RP-SME-001"
    )
    assert target.disposition.value == "manual_confirmation"
    assert target.included_in_formal is False
    assert "LLM二审" in target.rationale


def test_default_enhanced_mode_prefers_high_value_llm_tasks() -> None:
    text = """
    项目属性：服务
    评分方法：综合评分法
    方案评分扣分模式：完全满足/不完全满足。
    本项目专门面向中小企业采购，仍适用价格扣除。
    """
    base_report = TenderReviewEngine(review_mode=ReviewMode.fast).review_text(
        text, document_name="demo.txt"
    )
    enhanced_report = QwenReviewEnhancer(client=FakeClient()).enhance(base_report)
    llm_tasks = {item.task_name: item.status.value for item in enhanced_report.task_records if item.task_name.startswith("llm_")}
    assert llm_tasks["llm_scenario_review"] == "completed"
    assert llm_tasks["llm_scoring_review"] == "completed"
    assert llm_tasks["llm_review_point_second_review"] == "completed"
    assert llm_tasks["llm_clause_supplement"] == "skipped"
    assert llm_tasks["llm_specialist_review"] == "skipped"
    assert llm_tasks["llm_consistency_review"] == "skipped"


def test_review_point_second_review_prompt_includes_intensity_judgment_guidance() -> None:
    text = """
    预算金额：578600.00元
    评分方法：综合评分法
    商务部分 | 资质证书 (5.0分) | 每具有1个相关认证证书得1分。
    商务部分 | 管理体系认证情况 (5.0分) | 质量管理体系认证得分。
    """
    report = TenderReviewEngine(review_mode=ReviewMode.fast).review_text(text, document_name="demo.txt")
    prompt = build_review_point_second_review_prompt(report)
    assert "intensity_judgment" in prompt
    assert "primary_evidence_judgment" in prompt
    assert "supporting_evidence_judgment" in prompt
    assert "偏重要求" in prompt
    assert "刚性门槛" in prompt
    assert "裁量过大" in prompt


def test_dynamic_tasks_can_add_evidence_hints_rebuttal_templates_and_enhanced_assembly() -> None:
    text = """
    项目属性：货物
    采购标的：苗木供货，仅供货，不含人工服务。
    合同履行期限：1095日。
    造林内容包含人工管护、抚育和运水。
    """
    base_report = TenderReviewEngine(review_mode=ReviewMode.fast).review_text(
        text, document_name="demo.txt"
    )
    enhanced_report = QwenReviewEnhancer(client=FakeClient()).enhance(base_report)

    dynamic_point = next(item for item in enhanced_report.review_points if item.title == "项目属性与采购内容结构错配")
    assert any("动态任务补证提示" in note for note in dynamic_point.evidence_bundle.missing_evidence_notes)
    assert dynamic_point.evidence_bundle.rebuttal_evidence
    assert any(
        "仅供货" in item.quote or "不含人工服务" in item.quote
        for item in dynamic_point.evidence_bundle.rebuttal_evidence
    )
    assert any(
        "项目属性=货物" in item.quote
        for item in dynamic_point.evidence_bundle.direct_evidence + dynamic_point.evidence_bundle.supporting_evidence
    )
    assert "专属组证增强" in dynamic_point.rationale
    assert "structure 类型执行差异化组证" in dynamic_point.rationale


def test_dynamic_tasks_can_use_type_specific_scoring_assembly() -> None:
    text = """
    项目属性：服务
    评分方法：综合评分法
    财务指标：营业收入越高得分越高。
    样品分：10分
    """
    dynamic_definition = [
        {
            "catalog_id": "RP-DYN-002",
            "title": "评分项与项目履约能力相关性复核",
            "dimension": "评审标准明确性",
            "severity": "high",
            "task_type": "scoring",
            "scenario_tags": ["dynamic", "scoring"],
            "focus_fields": ["评分方法"],
            "signal_groups": [["营业收入", "得分"]],
            "evidence_hints": ["优先采集评分方法、财务指标和样品分条款"],
            "rebuttal_templates": [],
            "enhancement_fields": ["评分方法"],
            "basis_hint": "评分相关动态任务应优先围绕评分字段组证。",
        }
    ]
    base_report = TenderReviewEngine(review_mode=ReviewMode.fast).review_text(
        text, document_name="demo.txt"
    )
    from agent_review.llm.task_planner import build_dynamic_review_points, parse_dynamic_review_tasks

    points = build_dynamic_review_points(
        parse_dynamic_review_tasks(dynamic_definition),
        base_report.extracted_clauses,
    )
    scoring_point = points[0]

    quotes = [item.quote for item in scoring_point.evidence_bundle.direct_evidence + scoring_point.evidence_bundle.supporting_evidence]
    assert any("评分方法" in quote for quote in quotes)
    assert any("财务指标" in quote or "营业收入" in quote for quote in quotes)
    assert any("样品分" in quote for quote in quotes)
    assert "scoring 类型执行差异化组证" in scoring_point.rationale


def test_dynamic_task_type_can_flow_into_llm_second_review_prompt_and_override() -> None:
    text = """
    项目属性：货物
    采购标的：苗木供货，仅供货，不含人工服务。
    合同履行期限：1095日。
    造林内容包含人工管护、抚育和运水。
    """
    base_report = TenderReviewEngine(review_mode=ReviewMode.fast).review_text(
        text, document_name="demo.txt"
    )
    scenario_report = QwenReviewEnhancer(client=FakeClient()).enhance(base_report)

    prompt = build_review_point_second_review_prompt(scenario_report)
    assert '"task_type": "structure"' in prompt
    assert "重点复核项目属性、采购内容、合同类型、履约周期之间是否真实错配" in prompt

    second_review_report = QwenReviewEnhancer(client=FakeDynamicTypedSecondReviewClient()).enhance(scenario_report)
    assert any(
        item.title == "项目属性与采购内容结构错配" and "结构类二审重点已触发" in item.rationale
        for item in second_review_report.llm_semantic_review.review_point_second_reviews
    )


def test_scoring_review_prompt_and_dynamic_tasks_can_flow_into_main_chain() -> None:
    text = """
    项目属性：货物
    采购标的：家具
    评分方法：综合评分法
    方案评分扣分模式：完全满足且优于/完全满足/不完全满足。
    行业相关性存疑评分项：软件企业认定证书5分、财务报告2分。
    """
    base_report = TenderReviewEngine(review_mode=ReviewMode.fast).review_text(
        text, document_name="demo.txt"
    )
    enhanced_report = QwenReviewEnhancer(client=FakeClient()).enhance(base_report)

    assert enhanced_report.llm_semantic_review.scoring_review_summary
    assert enhanced_report.llm_semantic_review.scoring_dynamic_review_tasks
    scoring_task = enhanced_report.llm_semantic_review.scoring_dynamic_review_tasks[0]
    assert scoring_task.task_type == "scoring"
    assert scoring_task.evidence_hints
    assert any(item.title == scoring_task.title for item in enhanced_report.review_points)
    assert any(
        item.title == "证书检测报告及财务指标权重合理性复核"
        for item in enhanced_report.review_points
    )


def test_dynamic_task_parser_can_enrich_defaults_for_unknown_task_shape() -> None:
    from agent_review.llm.task_planner import parse_dynamic_review_tasks

    parsed = parse_dynamic_review_tasks(
        [
            {
                "title": "需求调查结论与项目复杂度匹配性复核",
                "task_type": "generic",
            }
        ]
    )
    task = parsed[0]
    assert task.dimension == "综合风险复核"
    assert "需求调查结论" in [field for condition in task.required_conditions for field in condition.clause_fields]
    assert task.evidence_hints
    assert task.enhancement_fields


def test_extractors_can_capture_demand_survey_and_scoring_item_details() -> None:
    text = """
    本项目不需要需求调查。
    不组织专家论证。
    项目整体实施方案（3分）：完全满足且优于项目要求的得3分。
    售后服务方案（3分）：完全满足项目要求的得2分。
    """
    clauses = extract_clauses(text)
    clause_map = {item.field_name: item for item in clauses}
    assert clause_map["需求调查结论"].normalized_value == "不需要"
    assert clause_map["专家论证结论"].normalized_value == "不需要"
    assert "项目整体实施方案" in clause_map["评分项明细"].content
    assert "售后服务方案" in clause_map["评分项明细"].content


def test_scoring_weight_point_can_distinguish_bid_stage_submission() -> None:
    text = """
    项目属性：货物
    采购标的：家具
    评分方法：综合评分法
    软件企业认定证书5分，财务报告2分。
    投标文件中提供认证证书和检测报告，作为评分依据。
    """
    clauses = extract_clauses(text)
    point = ReviewPoint(
        point_id="RP-T-001",
        catalog_id="RP-SCORE-008",
        title="证书检测报告及财务指标权重合理性复核",
        dimension="评审标准明确性",
        severity=Severity.high,
        status=ReviewPointStatus.identified,
        rationale="测试评分材料适用阶段。",
        evidence_bundle=EvidenceBundle(),
        source_findings=["task_library:RP-SCORE-008"],
    )
    check = build_applicability_checks([point], clauses)[0]
    assert check.applicable is True
    assert "存在证书报告或财务指标评分信号" in check.satisfied_conditions
    assert "投标阶段" in check.summary or any("投标阶段" in item.detail for item in check.requirement_results)


def test_scoring_weight_point_can_exclude_delivery_stage_materials() -> None:
    text = """
    项目属性：货物
    采购标的：家具
    评分方法：综合评分法
    软件企业认定证书5分。
    中标后供货验收时提供认证证书和检测报告。
    """
    clauses = extract_clauses(text)
    point = ReviewPoint(
        point_id="RP-T-002",
        catalog_id="RP-SCORE-008",
        title="证书检测报告及财务指标权重合理性复核",
        dimension="评审标准明确性",
        severity=Severity.high,
        status=ReviewPointStatus.identified,
        rationale="测试履约阶段材料排除。",
        evidence_bundle=EvidenceBundle(),
        source_findings=["task_library:RP-SCORE-008"],
    )
    check = build_applicability_checks([point], clauses)[0]
    excluded = next(item for item in check.exclusion_results if item.name == "证书检测报告仅在履约或验收阶段提交")
    assert excluded.status == ApplicabilityStatus.excluded
    assert check.applicable is False
    assert "要件链被阻断" in check.summary


def test_rigid_patent_requirement_can_be_formally_supported() -> None:
    text = """
    采购标的：家具
    投标人必须具备与采购标的相关的外观、结构、工艺及技术专利。
    """
    clauses = extract_clauses(text)
    patent_clause = next(item for item in clauses if item.field_name == "是否要求专利")
    assert patent_clause.normalized_value == "刚性门槛"

    point = ReviewPoint(
        point_id="RP-T-003",
        catalog_id="RP-REST-004",
        title="刚性门槛型专利要求",
        dimension="A.限制竞争风险",
        severity=Severity.high,
        status=ReviewPointStatus.identified,
        rationale="测试专利刚性门槛。",
        evidence_bundle=EvidenceBundle(),
        source_findings=["task_library:RP-REST-004"],
    )
    check = build_applicability_checks([point], clauses)[0]
    assert check.applicable is True
    assert any("刚性门槛" in item.detail for item in check.requirement_results)


def test_rigid_patent_requirement_still_extracts_when_line_also_mentions_infringement() -> None:
    text = """
    中标人必须具备与采购标的相关的外观、结构、工艺及技术专利（专利须在保护期内），并保证不侵犯第三方知识产权。
    """
    clauses = extract_clauses(text)
    patent_clause = next(item for item in clauses if item.field_name == "是否要求专利")
    assert patent_clause.normalized_value == "刚性门槛"


def test_bid_stage_material_burden_point_prefers_bid_submission_context() -> None:
    text = """
    采购标的：家具
    投标文件中须提供主要原材料检测报告和质量管理体系认证证书作为评审依据。
    """
    clauses = extract_clauses(text)
    point = ReviewPoint(
        point_id="RP-T-004",
        catalog_id="RP-SCORE-009",
        title="投标阶段证书或检测报告负担过重",
        dimension="评审标准明确性",
        severity=Severity.high,
        status=ReviewPointStatus.identified,
        rationale="测试投标阶段材料负担。",
        evidence_bundle=EvidenceBundle(),
        source_findings=["task_library:RP-SCORE-009"],
    )
    check = build_applicability_checks([point], clauses)[0]
    assert check.applicable is True
    assert any("投标阶段" in item.detail for item in check.requirement_results)


def test_certificate_score_weight_point_uses_total_score() -> None:
    text = """
    预算金额：578600.00元
    商务部分 | 资质证书 (5.0分) | 每具有1个相关认证证书得1分。
    商务部分 | 管理体系认证情况 (5.0分) | 质量管理体系认证得分。
    """
    clauses = extract_clauses(text)
    score_clause = next(item for item in clauses if item.field_name == "证书类评分总分")
    assert score_clause.normalized_value == "10.0"

    point = ReviewPoint(
        point_id="RP-T-005",
        catalog_id="RP-SCORE-010",
        title="证书类评分分值偏高",
        dimension="评审标准明确性",
        severity=Severity.high,
        status=ReviewPointStatus.identified,
        rationale="测试证书类评分总分。",
        evidence_bundle=EvidenceBundle(),
        source_findings=["task_library:RP-SCORE-010"],
    )
    check = build_applicability_checks([point], clauses)[0]
    assert check.applicable is True
    assert any("10.0分" in item.detail for item in check.requirement_results)


def test_contract_template_residue_point_detects_placeholder_terms() -> None:
    text = """
    合同签订之日起1个月内完成设计、测试、验收。
    提供X年的免费质保服务。
    于事件发生后天内完成处理。
    """
    clauses = extract_clauses(text)
    point = ReviewPoint(
        point_id="RP-T-006",
        catalog_id="RP-TPL-007",
        title="合同文本存在明显模板残留",
        dimension="模板残留与冲突风险",
        severity=Severity.high,
        status=ReviewPointStatus.identified,
        rationale="测试合同模板残留。",
        evidence_bundle=EvidenceBundle(),
        source_findings=["task_library:RP-TPL-007"],
    )
    check = build_applicability_checks([point], clauses)[0]
    assert check.applicable is True
    assert any("合同模板残留" in item.detail for item in check.requirement_results)


def test_scoring_quant_clause_aggregates_multiple_similar_scoring_evidences() -> None:
    text = """
    评分方法：综合评分法
    项目整体实施方案：完全满足且优于项目要求的得10分，完全满足项目要求的得6分，不完全满足项目要求的得2分。
    售后服务方案：完全满足且优于项目要求的得8分，完全满足项目要求的得4分，不完全满足项目要求的得1分。
    """
    clauses = extract_clauses(text)
    scoring_clause = next(item for item in clauses if item.field_name == "方案评分扣分模式")

    assert "项目整体实施方案" in scoring_clause.content
    assert "售后服务方案" in scoring_clause.content
    assert "完全满足且优于" in scoring_clause.content
    assert "不完全满足项目要求" in scoring_clause.content


def test_material_burden_clause_aggregates_multiple_bid_stage_requirements() -> None:
    text = """
    投标文件中须提供主要原材料检测报告。
    投标文件中须提供质量管理体系认证证书、中国环保产品认证、中国环境标志产品认证。
    """
    clauses = extract_clauses(text)
    burden_clause = next(item for item in clauses if item.field_name == "证书检测报告负担特征")

    assert "检测报告" in burden_clause.content
    assert "质量管理体系认证证书" in burden_clause.content
    assert "中国环境标志产品认证" in burden_clause.content


def test_risk_rules_do_not_treat_factory_quality_wording_as_origin_brand_restriction() -> None:
    text = """
    中标人提供的货物是全新、表面和内部均无瑕疵的原厂正品。
    产品符合国家现行相关标准和厂家出厂标准。
    """
    hits = match_risk_rules(text)
    assert all(item.rule_name != "指定品牌/原厂限制" for item in hits)
    assert all(item.rule_name != "产地厂家商标限制" for item in hits)


def test_risk_rules_do_not_treat_same_brand_rule_or_origin_proof_as_restrictive() -> None:
    text = """
    提供相同品牌产品的，按评标办法处理。
    进口设备必须具备有效的原产地证明。
    """
    hits = match_risk_rules(text)
    assert all(item.rule_name != "指定品牌/原厂限制" for item in hits)
    assert all(item.rule_name != "产地厂家商标限制" for item in hits)


def test_contract_type_extractor_ignores_generic_sales_or_service_contract_wording() -> None:
    text = """
    根据上述业绩情况，按招标文件要求附销售或服务合同复印件及评审标准要求的证明材料。
    """
    clauses = extract_clauses(text)
    assert all(item.field_name != "合同类型" for item in clauses)


def test_formal_filters_generic_patent_when_rigid_patent_exists() -> None:
    generic = ReviewPoint(
        point_id="RP-G",
        catalog_id="RP-REST-003",
        title="专利要求",
        dimension="A.限制竞争风险",
        severity=Severity.high,
        status=ReviewPointStatus.suspected,
        rationale="泛化专利要求。",
        evidence_bundle=EvidenceBundle(
            direct_evidence=[Evidence(quote="是否要求专利=刚性门槛", section_hint="line:10")],
        ),
        source_findings=["task_library:RP-REST-003"],
        legal_basis=[LegalBasis(source_name="测试依据", article_hint="第1条", summary="测试用")],
    )
    rigid = ReviewPoint(
        point_id="RP-R",
        catalog_id="RP-REST-004",
        title="刚性门槛型专利要求",
        dimension="A.限制竞争风险",
        severity=Severity.high,
        status=ReviewPointStatus.suspected,
        rationale="刚性门槛型专利要求。",
        evidence_bundle=EvidenceBundle(
            direct_evidence=[Evidence(quote="是否要求专利=刚性门槛", section_hint="line:10")],
        ),
        source_findings=["task_library:RP-REST-004"],
        legal_basis=[LegalBasis(source_name="测试依据", article_hint="第1条", summary="测试用")],
    )
    checks = [
        ApplicabilityCheck(
            point_id="RP-G",
            catalog_id="RP-REST-003",
            applicable=True,
            requirement_results=[],
            exclusion_results=[],
            satisfied_conditions=["存在专利要求"],
            missing_conditions=[],
            blocking_conditions=[],
            requirement_chain_complete=True,
            summary="要件链成立。",
        ),
        ApplicabilityCheck(
            point_id="RP-R",
            catalog_id="RP-REST-004",
            applicable=True,
            requirement_results=[],
            exclusion_results=[],
            satisfied_conditions=["专利要求具有刚性门槛特征"],
            missing_conditions=[],
            blocking_conditions=[],
            requirement_chain_complete=True,
            summary="要件链成立。",
        ),
    ]
    gates = [
        ReviewQualityGate(point_id="RP-G", status=QualityGateStatus.passed, reasons=[]),
        ReviewQualityGate(point_id="RP-R", status=QualityGateStatus.passed, reasons=[]),
    ]
    adjudications = build_formal_adjudication(
        [generic, rigid],
        checks,
        gates,
        "第10行：投标人必须具备相关专利。",
        [],
    )
    mapping = {item.catalog_id: item for item in adjudications}
    assert mapping["RP-REST-003"].included_in_formal is False
    assert mapping["RP-REST-004"].included_in_formal is True


def test_engine_can_review_multiple_files(tmp_path: Path) -> None:
    main_file = tmp_path / "main.txt"
    contract_file = tmp_path / "contract.txt"
    main_file.write_text("项目属性：服务\n采购需求：物业服务。", encoding="utf-8")
    contract_file.write_text("合同条款\n付款方式：按月支付。", encoding="utf-8")
    engine = TenderReviewEngine(review_mode=ReviewMode.fast)
    report = engine.review_files(
        [
            main_file,
            contract_file,
        ]
    )
    assert report.source_documents
    assert len(report.source_documents) == 2
    markdown = render_markdown(report)
    assert "## 联合审查文件" in markdown


def test_multi_file_review_detects_cross_document_consistency_issues(tmp_path: Path) -> None:
    tender = tmp_path / "tender.txt"
    scoring = tmp_path / "scoring.txt"
    contract = tmp_path / "contract.txt"
    tender.write_text(
        "\n".join(
            [
                "投标邀请",
                "采购需求",
                "项目属性：服务",
                "本项目专门面向中小企业采购。",
            ]
        ),
        encoding="utf-8",
    )
    scoring.write_text(
        "\n".join(
            [
                "评分标准",
                "综合评分",
                "价格扣除按10%执行。",
            ]
        ),
        encoding="utf-8",
    )
    contract.write_text(
        "\n".join(
            [
                "合同条款",
                "付款方式：尾款支付以采购人满意度考核为准。",
                "验收条款：按合同约定执行。",
            ]
        ),
        encoding="utf-8",
    )

    report = TenderReviewEngine(review_mode=ReviewMode.fast).review_files(
        [tender, scoring, contract]
    )

    titles = {item.title for item in report.findings}
    assert "正文 vs 评分细则跨文件一致性" in titles
    assert "正文 vs 合同草案跨文件一致性" in titles
    markdown = render_markdown(report)
    assert "## 跨文件一致性专项" in markdown


def test_fast_mode_skips_enhancer() -> None:
    text = "项目概况\n采购需求详见附件。"
    engine = TenderReviewEngine(review_enhancer=FakeEnhancer(), review_mode=ReviewMode.fast)
    report = engine.review_text(text, document_name="demo.txt")

    assert report.review_mode == ReviewMode.fast
    assert report.llm_enhanced is False


def test_write_review_artifacts_outputs_base_and_final(tmp_path: Path) -> None:
    text = "项目概况\n采购需求详见附件。"
    base_engine = TenderReviewEngine(review_mode=ReviewMode.fast)
    base_report = base_engine.review_text(text, document_name="demo.txt")

    enhanced_engine = TenderReviewEngine(review_enhancer=FakeEnhancer(), review_mode=ReviewMode.enhanced)
    enhanced_report = enhanced_engine.review_text(text, document_name="demo.txt")

    bundle = write_review_artifacts(enhanced_report, base_report, tmp_path)

    assert Path(bundle.base_json_path).exists()
    assert Path(bundle.base_markdown_path).exists()
    assert Path(bundle.final_json_path).exists()
    assert Path(bundle.final_markdown_path).exists()
    assert Path(bundle.opinion_letter_path).exists()
    assert Path(bundle.formal_review_opinion_path).exists()
    assert Path(bundle.manifest_path).exists()
    assert Path(bundle.llm_tasks_path).exists()
    assert Path(bundle.high_risk_review_path).exists()
    assert Path(bundle.pending_confirmation_path).exists()
    assert Path(bundle.specialist_table_paths["sme_policy"]["base"]).exists()
    assert Path(bundle.specialist_table_paths["sme_policy"]["final"]).exists()

    manifest = json.loads(Path(bundle.manifest_path).read_text(encoding="utf-8"))
    assert manifest["artifact_paths"]["base_report"]["json"] == bundle.base_json_path
    assert manifest["artifact_paths"]["final_report"]["json"] == bundle.final_json_path
    assert manifest["artifact_paths"]["opinion_letter"] == bundle.opinion_letter_path
    assert manifest["artifact_paths"]["formal_review_opinion"] == bundle.formal_review_opinion_path
    assert manifest["stage_records"]
    assert "core_modules" in manifest["rule_selection"]
    assert "enhancement_modules" in manifest["rule_selection"]
    llm_tasks_payload = json.loads(Path(bundle.llm_tasks_path).read_text(encoding="utf-8"))
    assert all(item["task_name"].startswith("llm_") for item in llm_tasks_payload["tasks"])
    pending_payload = json.loads(Path(bundle.pending_confirmation_path).read_text(encoding="utf-8"))
    assert "items" in pending_payload


def test_write_review_artifacts_outputs_specialist_table_files(tmp_path: Path) -> None:
    text = """
    项目属性：服务
    中小企业声明函：制造商声明
    本项目专门面向中小企业采购，仍适用价格扣除。
    年龄要求35岁以下。
    采购人拥有最终解释权。
    """
    base_report = TenderReviewEngine(review_mode=ReviewMode.fast).review_text(text, document_name="demo.txt")
    enhanced_report = TenderReviewEngine(
        review_enhancer=FakeEnhancer(),
        review_mode=ReviewMode.enhanced,
    ).review_text(text, document_name="demo.txt")

    bundle = write_review_artifacts(enhanced_report, base_report, tmp_path)

    for table_name in [
        "project_structure",
        "sme_policy",
        "personnel_boundary",
        "contract_performance",
        "template_conflicts",
    ]:
        assert Path(bundle.specialist_table_paths[table_name]["base"]).exists()
        assert Path(bundle.specialist_table_paths[table_name]["final"]).exists()


def test_markdown_report_uses_v2_sections() -> None:
    text = """
    项目属性：服务
    本项目专门面向中小企业采购，仍适用价格扣除。
    年龄要求35岁以下。
    采购人拥有最终解释权。
    """
    report = TenderReviewEngine().review_text(text, document_name="demo.txt")
    markdown = render_markdown(report)

    assert "## 高风险问题" in markdown
    assert "## 中风险问题" in markdown
    assert "## 审查边界说明" in markdown
    assert "## 中小企业政策一致性表" in markdown


def test_markdown_can_render_legal_basis() -> None:
    text = """
    采购需求
    本项目要求指定品牌产品。
    """
    report = TenderReviewEngine().review_text(text, document_name="demo.txt")
    markdown = render_markdown(report)

    assert "法规依据" in markdown
    assert "中华人民共和国政府采购法" in markdown


def test_opinion_letter_can_render_formal_sections() -> None:
    text = """
    项目属性：服务
    本项目专门面向中小企业采购，仍适用价格扣除。
    年龄要求35岁以下。
    采购人拥有最终解释权。
    """
    report = TenderReviewEngine(review_mode=ReviewMode.fast).review_text(text, document_name="demo.txt")

    opinion = render_opinion_letter(report)

    assert "# 招标文件审查意见书" in opinion
    assert "## 三、审查结论" in opinion
    assert "## 四、主要审查意见" in opinion
    assert "事实要件：" in opinion
    assert "## 五、修改建议" in opinion
    assert "## 七、审查边界说明" in opinion


def test_formal_review_opinion_can_render_high_risk_fields() -> None:
    text = """
    项目属性：服务
    本项目专门面向中小企业采购，仍适用价格扣除。
    年龄要求35岁以下。
    采购人拥有最终解释权。
    """
    report = TenderReviewEngine(review_mode=ReviewMode.fast).review_text(text, document_name="demo.txt")

    formal = render_formal_review_opinion(report)

    assert "# 招标文件高风险正式审查意见" in formal
    assert "- 问题标题:" in formal
    assert "- 条款位置:" in formal
    assert "- 原文摘录:" in formal
    assert "- 问题类型:" in formal
    assert "- 风险等级:" in formal
    assert "- 合规判断:" in formal
    assert "- 法律/政策依据:" in formal


def test_formal_review_opinion_filters_template_and_weak_hits() -> None:
    text = """
    法定代表人证明书
    附：代表人性别：_____年龄：_________ 身份证号码：__________________
    一、名词解释
    采购代理机构：本项目是指某采购中心，对招标文件拥有最终的解释权。
    本项目属于专门面向中小企业采购的项目。
    对小型、微型企业给予价格扣除。
    """
    report = TenderReviewEngine(review_mode=ReviewMode.fast).review_text(text, document_name="demo.txt")

    formal = render_formal_review_opinion(report)

    assert "性别限制" not in formal
    assert "年龄限制" not in formal
    assert "采购人单方解释或决定条款" not in formal
    assert "价格扣除" in formal
    family_items = [
        item
        for item in report.formal_adjudication
        if "价格扣除" in item.title and "中小企业" in item.title
    ]
    assert family_items
    assert any(item.included_in_formal for item in family_items)
    assert any(item.evidence_sufficient for item in family_items if item.included_in_formal)
    assert any(item.legal_basis_applicable for item in family_items if item.included_in_formal)


def test_formal_review_opinion_can_render_review_layer_for_manual_confirmation() -> None:
    text = """
    项目属性：货物
    采购标的：复合型项目
    合同履行期限：1095日。
    本项目不需要需求调查。
    人工管护、抚育、运水等作业内容由供应商负责。
    """
    report = TenderReviewEngine(review_mode=ReviewMode.fast).review_text(text, document_name="demo.txt")
    formal = render_formal_review_opinion(report)

    assert "## 建议复核问题" in formal


def test_formal_review_opinion_suppresses_review_mirror_items() -> None:
    text = """
    项目属性：货物
    评分方法：综合评分法
    项目整体实施方案：完全满足且优于/完全满足/不完全满足。
    售后服务方案：完全满足且优于/完全满足/不完全满足。
    """
    report = TenderReviewEngine(review_mode=ReviewMode.fast).review_text(text, document_name="demo.txt")
    report.review_points.append(
        ReviewPoint(
            point_id="RP-MANUAL-001",
            catalog_id="RP-SCORE-007",
            title="评分分档主观性与量化充分性复核",
            dimension="评审标准明确性",
            severity=Severity.high,
            status=ReviewPointStatus.manual_confirmation,
            rationale="当前评分分档与 formal 已覆盖问题存在镜像关系。",
            evidence_bundle=EvidenceBundle(
                supporting_evidence=[Evidence(quote="完全满足/不完全满足", section_hint="line:4")],
                sufficiency_summary="主证据仍需压缩。",
            ),
            source_findings=["task_library:RP-SCORE-007"],
        )
    )
    report.formal_adjudication.append(
        FormalAdjudication(
            point_id="RP-MANUAL-001",
            catalog_id="RP-SCORE-007",
            title="评分分档主观性与量化充分性复核",
            disposition=FormalDisposition.manual_confirmation,
            rationale="当前与 formal 已有评分量化不足存在镜像重复。",
            included_in_formal=False,
            section_hint="line:4",
            primary_quote="当前自动抽取未定位到可直接引用的原文。",
            evidence_sufficient=False,
            legal_basis_applicable=False,
            applicability_summary="要件链未闭合。",
            quality_gate_status=QualityGateStatus.manual_confirmation,
            recommended_for_review=True,
            review_reason="镜像重复，且主证据代表性不足。",
        )
    )

    formal = render_formal_review_opinion(report)

    assert "方案评分量化不足" in formal
    assert "评分分档主观性与量化充分性复核" not in formal


def test_prudential_review_points_can_enter_review_layer() -> None:
    text = """
    项目属性：货物
    采购标的：苗木、肥料、防治药剂、标识牌及三年管护
    合同履行期限：1095日
    本项目不需要需求调查。
    不组织专家论证。
    """
    report = TenderReviewEngine(review_mode=ReviewMode.fast).review_text(text, document_name="demo.txt")
    formal = render_formal_review_opinion(report)

    titles = {item.title for item in report.review_points}
    assert "需求调查结论与项目复杂度匹配性复核" in titles
    assert "专家论证必要性建议复核" in titles
    applicability_titles = {item.catalog_id for item in report.applicability_checks if item.catalog_id.startswith("RP-PRUD-")}
    assert "RP-PRUD-001" in applicability_titles
    assert "RP-PRUD-002" in applicability_titles


def test_report_contains_specialist_tables() -> None:
    text = """
    项目属性：服务
    中小企业声明函：制造商声明
    本项目专门面向中小企业采购，仍适用价格扣除。
    年龄要求35岁以下。
    采购人拥有最终解释权。
    """
    report = TenderReviewEngine().review_text(text, document_name="demo.txt")

    assert report.specialist_tables.sme_policy
    assert report.specialist_tables.personnel_boundary
    assert report.specialist_tables.contract_performance


def test_report_contains_review_point_and_formal_adjudication_skeleton() -> None:
    text = """
    项目属性：服务
    本项目专门面向中小企业采购，仍适用价格扣除。
    年龄要求35岁以下。
    """
    report = TenderReviewEngine().review_text(text, document_name="demo.txt")

    assert report.review_points
    assert report.formal_adjudication
    markdown = render_markdown(report)
    assert "## ReviewPoint" in markdown
    assert "## Formal Adjudication" in markdown


def test_rule_and_consistency_layers_prioritize_review_points() -> None:
    text = """
    项目属性：服务
    本项目专门面向中小企业采购，仍适用价格扣除。
    合同条款：验收合格后付款，但尾款根据采购人满意度考核后支付。
    """
    report = TenderReviewEngine().review_text(text, document_name="demo.txt")

    point_map = {(item.dimension, item.title): item for item in report.review_points}

    rule_point = point_map[("中小企业政策风险", "专门面向中小企业却仍保留价格扣除")]
    assert any(source.startswith("risk_hit:") for source in rule_point.source_findings)
    assert rule_point.status == ReviewPointStatus.confirmed
    assert rule_point.evidence_bundle.direct_evidence

    consistency_point = point_map[("跨条款一致性检查", "验收标准 vs 付款条件")]
    assert any(source.startswith("consistency_check:") for source in consistency_point.source_findings)
    assert consistency_point.status == ReviewPointStatus.suspected
    assert consistency_point.evidence_bundle.missing_evidence_notes

    findings = {(item.dimension, item.title): item for item in report.findings}
    assert findings[("中小企业政策风险", "专门面向中小企业却仍保留价格扣除")].finding_type == FindingType.confirmed_issue
    assert findings[("跨条款一致性检查", "验收标准 vs 付款条件")].finding_type == FindingType.warning


def test_markdown_can_render_specialist_summary() -> None:
    text = """
    项目属性：服务
    中小企业声明函：制造商声明
    本项目专门面向中小企业采购，仍适用价格扣除。
    """
    engine = TenderReviewEngine(review_enhancer=FakeEnhancer(), review_mode=ReviewMode.enhanced)
    report = engine.review_text(text, document_name="demo.txt")
    markdown = render_markdown(report)

    assert "中小企业政策一致性表摘要" in markdown
    assert "这是经过LLM增强的中小企业政策专项摘要。" in markdown

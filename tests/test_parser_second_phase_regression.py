from pathlib import Path

from docx import Document

from agent_review.ontology import EffectTag, SemanticZoneType
from agent_review.parsers import load_document


def _build_second_phase_doc(file_path: Path) -> None:
    document = Document()
    document.add_paragraph("第一章 招标公告", style="Heading 1")
    document.add_paragraph("投标人资格要求", style="Heading 2")
    document.add_paragraph("投标人须具备相关资质。")

    document.add_paragraph("第二章 采购需求", style="Heading 1")
    document.add_paragraph("技术要求", style="Heading 2")
    document.add_paragraph("木质家具甲醛释放量应符合国家标准。")
    document.add_paragraph("商务要求", style="Heading 2")
    document.add_paragraph("交货期和售后服务应满足项目要求。")

    document.add_paragraph("第三章 评分办法", style="Heading 1")
    table = document.add_table(rows=2, cols=3)
    table.rows[0].cells[0].text = "评分项"
    table.rows[0].cells[1].text = "分值"
    table.rows[0].cells[2].text = "评分标准"
    table.rows[1].cells[0].text = "检测报告"
    table.rows[1].cells[1].text = "5"
    table.rows[1].cells[2].text = "提供得分"

    document.add_paragraph("第四章 投标文件格式、附件", style="Heading 1")
    document.add_paragraph("中小企业声明函（格式）")
    document.add_paragraph("采购需求详见附件1。")
    document.save(file_path)


def test_second_phase_parser_separates_qualification_technical_business_scoring_and_template(tmp_path: Path) -> None:
    file_path = tmp_path / "second_phase.docx"
    _build_second_phase_doc(file_path)

    _, parse_result = load_document(file_path)
    node_map = {item.node_id: item for item in parse_result.document_nodes}
    zone_map = {item.node_id: item for item in parse_result.semantic_zones}
    effect_map = {item.node_id: item for item in parse_result.effect_tag_results}

    qualification_node = next(item for item in parse_result.document_nodes if item.title == "投标人资格要求")
    technical_node = next(item for item in parse_result.document_nodes if item.title == "技术要求")
    business_node = next(item for item in parse_result.document_nodes if item.title == "商务要求")
    template_node = next(item for item in parse_result.document_nodes if "中小企业声明函（格式）" in item.title)
    appendix_node = next(item for item in parse_result.document_nodes if "详见附件1" in item.title)
    scoring_row = next(
        item
        for item in parse_result.document_nodes
        if item.node_type.value == "table_row" and "检测报告" in item.title
    )

    assert zone_map[qualification_node.node_id].zone_type == SemanticZoneType.qualification
    assert zone_map[technical_node.node_id].zone_type == SemanticZoneType.technical
    assert zone_map[business_node.node_id].zone_type == SemanticZoneType.business
    assert zone_map[template_node.node_id].zone_type == SemanticZoneType.template
    assert zone_map[appendix_node.node_id].zone_type == SemanticZoneType.appendix_reference
    assert zone_map[scoring_row.node_id].zone_type == SemanticZoneType.scoring

    assert EffectTag.binding in effect_map[qualification_node.node_id].effect_tags
    assert EffectTag.binding in effect_map[technical_node.node_id].effect_tags
    assert EffectTag.binding in effect_map[business_node.node_id].effect_tags
    assert EffectTag.template in effect_map[template_node.node_id].effect_tags
    assert EffectTag.reference_only in effect_map[appendix_node.node_id].effect_tags
    assert any(unit.zone_type == SemanticZoneType.scoring for unit in parse_result.clause_units)
    assert any(unit.zone_type == SemanticZoneType.template for unit in parse_result.clause_units)
    assert any(unit.zone_type == SemanticZoneType.appendix_reference for unit in parse_result.clause_units)
    assert any(
        unit.zone_type == SemanticZoneType.scoring and "检测报告" in unit.text
        for unit in parse_result.clause_units
    )

    assert "第二章 采购需求" in node_map[technical_node.node_id].path
    assert node_map[scoring_row.node_id].path


def test_second_phase_clause_units_keep_weak_effect_text_out_of_binding_chain(tmp_path: Path) -> None:
    file_path = tmp_path / "weak_effect.docx"
    document = Document()
    document.add_paragraph("第三章 投标文件格式、附件", style="Heading 1")
    document.add_paragraph("中小企业声明函（格式）")
    document.add_paragraph("详见附件2。")
    document.save(file_path)

    _, parse_result = load_document(file_path)

    weak_units = [
        unit
        for unit in parse_result.clause_units
        if unit.zone_type in {SemanticZoneType.template, SemanticZoneType.appendix_reference}
    ]

    assert weak_units
    assert any(EffectTag.template in unit.effect_tags for unit in weak_units if unit.zone_type == SemanticZoneType.template)
    assert any(EffectTag.reference_only in unit.effect_tags for unit in weak_units if unit.zone_type == SemanticZoneType.appendix_reference)

from pathlib import Path

from docx import Document

from agent_review.parsers import load_document
from agent_review.ontology import EffectTag


def _effect_map(parse_result):
    return {item.node_id: item for item in parse_result.effect_tag_results}


def test_effect_tagger_marks_template_example_optional_and_reference(tmp_path: Path) -> None:
    file_path = tmp_path / "effects.docx"
    document = Document()
    document.add_paragraph("第三章 投标文件格式、附件", style="Heading 1")
    document.add_paragraph("中小企业声明函（格式）", style="Heading 2")
    document.add_paragraph("以下为示例文本，可选填写，详见附件。")
    document.save(file_path)

    _, parse_result = load_document(file_path)
    effects = _effect_map(parse_result)
    node_map = {item.node_id: item for item in parse_result.document_nodes}

    title_node = next(item for item in parse_result.document_nodes if item.title == "中小企业声明函（格式）")
    text_node = next(item for item in parse_result.document_nodes if "以下为示例文本" in item.text)

    assert EffectTag.template in effects[title_node.node_id].effect_tags
    assert EffectTag.template in effects[text_node.node_id].effect_tags
    assert EffectTag.example in effects[text_node.node_id].effect_tags
    assert EffectTag.optional in effects[text_node.node_id].effect_tags
    assert EffectTag.reference_only in effects[text_node.node_id].effect_tags
    assert "投标文件格式" in node_map[title_node.node_id].path


def test_effect_tagger_marks_catalog_and_binding(tmp_path: Path) -> None:
    file_path = tmp_path / "catalog_binding.docx"
    document = Document()
    document.add_paragraph("目录")
    document.add_paragraph("第一章 招标公告")
    document.add_paragraph("第一章 招标公告", style="Heading 1")
    document.add_paragraph("投标人资格要求：投标人须具备相关资质。")
    document.save(file_path)

    _, parse_result = load_document(file_path)
    effects = _effect_map(parse_result)

    catalog_node = next(item for item in parse_result.document_nodes if item.node_type.value == "catalog_entry")
    binding_node = next(item for item in parse_result.document_nodes if "投标人资格要求" in item.text and item.node_type.value == "paragraph")

    assert effects[catalog_node.node_id].effect_tags == [EffectTag.catalog]
    assert effects[binding_node.node_id].effect_tags == [EffectTag.binding]

from pathlib import Path

from docx import Document

from agent_review.extractors import extract_clauses_from_units
from agent_review.parsers import load_document


def test_extract_clauses_from_units_can_extract_structured_fields(tmp_path: Path) -> None:
    file_path = tmp_path / "extract_from_units.docx"
    document = Document()
    document.add_paragraph("关键信息", style="Heading 1")
    document.add_paragraph("项目属性：货物")
    document.add_paragraph("投标人资格要求", style="Heading 1")
    document.add_paragraph("投标人须具备相关资质。")
    document.add_paragraph("综合评分法评标信息", style="Heading 1")
    table = document.add_table(rows=2, cols=3)
    table.rows[0].cells[0].text = "评分项"
    table.rows[0].cells[1].text = "分值"
    table.rows[0].cells[2].text = "评分标准"
    table.rows[1].cells[0].text = "检测报告"
    table.rows[1].cells[1].text = "5"
    table.rows[1].cells[2].text = "提供得分"
    document.save(file_path)

    _, parse_result = load_document(file_path)
    clauses = extract_clauses_from_units(parse_result.clause_units)
    field_names = {item.field_name for item in clauses}

    assert "项目属性" in field_names
    assert any(item.source_anchor for item in clauses)

from pathlib import Path

from docx import Document

from agent_review.parsers import load_document
from agent_review.ontology import ClauseSemanticType, EffectTag, SemanticZoneType


def test_clause_unit_builder_creates_units_for_paragraphs_and_table_rows(tmp_path: Path) -> None:
    file_path = tmp_path / "clause_units.docx"
    document = Document()
    document.add_paragraph("第一章 招标公告", style="Heading 1")
    document.add_paragraph("投标人资格要求", style="Heading 2")
    document.add_paragraph("投标人须具备相关资质。")
    document.add_paragraph("综合评分法评标信息", style="Heading 1")
    table = document.add_table(rows=2, cols=3)
    table.rows[0].cells[0].text = "评分项"
    table.rows[0].cells[1].text = "分值"
    table.rows[0].cells[2].text = "评分标准"
    table.rows[1].cells[0].text = "检测报告"
    table.rows[1].cells[1].text = "5"
    table.rows[1].cells[2].text = "提供得分"
    document.save(file_path)

    _, parse_result = load_document(file_path)

    assert parse_result.clause_units
    assert any(unit.zone_type == SemanticZoneType.qualification for unit in parse_result.clause_units)
    assert any(unit.zone_type == SemanticZoneType.scoring for unit in parse_result.clause_units)
    assert any(unit.clause_semantic_type == ClauseSemanticType.scoring_rule for unit in parse_result.clause_units)


def test_clause_unit_builder_marks_template_units(tmp_path: Path) -> None:
    file_path = tmp_path / "template_units.docx"
    document = Document()
    document.add_paragraph("第三章 投标文件格式、附件", style="Heading 1")
    document.add_paragraph("中小企业声明函（格式）", style="Heading 2")
    document.add_paragraph("投标人应按以下格式填写并盖章。")
    document.save(file_path)

    _, parse_result = load_document(file_path)

    template_units = [unit for unit in parse_result.clause_units if unit.zone_type == SemanticZoneType.template]
    assert template_units
    assert any(EffectTag.template in unit.effect_tags for unit in template_units)
    assert any(unit.clause_semantic_type == ClauseSemanticType.declaration_template for unit in template_units)

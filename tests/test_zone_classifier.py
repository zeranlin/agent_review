from pathlib import Path

from docx import Document

from agent_review.parsers import load_document
from agent_review.ontology import SemanticZoneType


def _zone_map(parse_result):
    return {item.node_id: item for item in parse_result.semantic_zones}


def test_zone_classifier_identifies_template_and_scoring(tmp_path: Path) -> None:
    file_path = tmp_path / "zone.docx"
    document = Document()
    document.add_paragraph("第三章 投标文件格式、附件", style="Heading 1")
    document.add_paragraph("中小企业声明函（格式）", style="Heading 2")
    document.add_paragraph("投标人应按以下格式填写。")
    document.add_paragraph("综合评分法评标信息", style="Heading 1")
    table = document.add_table(rows=2, cols=3)
    table.rows[0].cells[0].text = "评分项"
    table.rows[0].cells[1].text = "分值"
    table.rows[0].cells[2].text = "评分标准"
    table.rows[1].cells[0].text = "检测报告"
    table.rows[1].cells[1].text = "5"
    table.rows[1].cells[2].text = "提供得分"
    document.save(file_path)

    _, parse_result = load_document(file_path)
    zones = _zone_map(parse_result)
    node_map = {item.node_id: item for item in parse_result.document_nodes}

    template_node = next(item for item in parse_result.document_nodes if "中小企业声明函（格式）" in item.title)
    scoring_table_row = next(item for item in parse_result.document_nodes if item.node_type.value == "table_row")

    assert zones[template_node.node_id].zone_type == SemanticZoneType.template
    assert zones[scoring_table_row.node_id].zone_type == SemanticZoneType.scoring
    assert "投标文件格式" in node_map[template_node.node_id].path


def test_zone_classifier_identifies_qualification_and_technical(tmp_path: Path) -> None:
    file_path = tmp_path / "zone2.docx"
    document = Document()
    document.add_paragraph("第一章 招标公告", style="Heading 1")
    document.add_paragraph("投标人资格要求", style="Heading 2")
    document.add_paragraph("投标人须具备相关资质。")
    document.add_paragraph("第二章 招标项目需求", style="Heading 1")
    document.add_paragraph("四、具体技术要求", style="Heading 2")
    document.add_paragraph("家具甲醛释放量应符合国家标准。")
    document.save(file_path)

    _, parse_result = load_document(file_path)
    zones = _zone_map(parse_result)

    qualification_node = next(item for item in parse_result.document_nodes if item.title == "投标人资格要求")
    technical_node = next(item for item in parse_result.document_nodes if item.title == "四、具体技术要求")

    assert zones[qualification_node.node_id].zone_type == SemanticZoneType.qualification
    assert zones[technical_node.node_id].zone_type == SemanticZoneType.technical

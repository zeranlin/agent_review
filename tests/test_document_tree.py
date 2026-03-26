from pathlib import Path

from docx import Document

from agent_review.parsers import load_document


def test_build_document_tree_generates_paths_and_catalog_entries(tmp_path: Path) -> None:
    file_path = tmp_path / "tree.docx"
    document = Document()
    document.add_paragraph("目录", style="Title")
    document.add_paragraph("第一章 招标公告")
    document.add_paragraph("第二章 招标项目需求")
    document.add_paragraph("第一章 招标公告", style="Heading 1")
    document.add_paragraph("投标人资格要求")
    document.add_paragraph("投标人须具备相关资质。")
    document.add_paragraph("第二章 招标项目需求", style="Heading 1")
    document.add_paragraph("技术要求")
    document.add_paragraph("提供符合标准的家具产品。")
    document.add_paragraph("商务要求")
    document.add_paragraph("交货期和售后服务应满足项目要求。")
    document.add_paragraph("第三章 投标文件格式、附件", style="Heading 1")
    document.add_paragraph("中小企业声明函（格式）")
    appendix_table = document.add_table(rows=2, cols=2)
    appendix_table.rows[0].cells[0].text = "条目"
    appendix_table.rows[0].cells[1].text = "说明"
    appendix_table.rows[1].cells[0].text = "报价表"
    appendix_table.rows[1].cells[1].text = "按格式填写"
    document.add_paragraph("采购需求详见附件1。")
    document.save(file_path)

    _, parse_result = load_document(file_path)

    assert parse_result.document_nodes
    catalog_nodes = [item for item in parse_result.document_nodes if item.node_type.value == "catalog_entry"]
    assert catalog_nodes
    assert any(item.title == "第一章 招标公告" for item in catalog_nodes)
    assert any(item.title == "第二章 招标项目需求" for item in catalog_nodes)
    qualification_node = next(item for item in parse_result.document_nodes if item.title == "投标人资格要求")
    technical_node = next(item for item in parse_result.document_nodes if item.title == "技术要求")
    business_node = next(item for item in parse_result.document_nodes if item.title == "商务要求")
    appendix_node = next(item for item in parse_result.document_nodes if item.title == "中小企业声明函（格式）")
    appendix_reference_node = next(item for item in parse_result.document_nodes if "详见附件1" in item.title)
    table_node = next(item for item in parse_result.document_nodes if item.node_type.value == "table")
    row_node = next(item for item in parse_result.document_nodes if item.node_type.value == "table_row")

    third_chapter = next(item for item in parse_result.document_nodes if item.title == "第三章 投标文件格式、附件")
    second_chapter = next(
        item
        for item in parse_result.document_nodes
        if item.title == "第二章 招标项目需求" and item.node_type.value == "chapter"
    )

    assert qualification_node.path.startswith("ROOT > 投标人资格要求")
    assert technical_node.parent_id == second_chapter.node_id
    assert business_node.parent_id == second_chapter.node_id
    assert business_node.parent_id != technical_node.node_id
    assert appendix_node.node_type.value == "appendix"
    assert appendix_node.parent_id == third_chapter.node_id
    assert appendix_reference_node.node_type.value == "appendix"
    assert "详见附件1" in appendix_reference_node.title
    assert table_node.parent_id == appendix_node.node_id
    assert row_node.parent_id == table_node.node_id
    assert table_node.path.startswith(appendix_node.path)
    assert row_node.path.startswith(f"{table_node.path} > row:1")


def test_document_tree_keeps_table_rows_as_nodes(tmp_path: Path) -> None:
    file_path = tmp_path / "table_tree.docx"
    document = Document()
    document.add_paragraph("综合评分法评标信息", style="Heading 1")
    table = document.add_table(rows=2, cols=3)
    table.rows[0].cells[0].text = "评分项"
    table.rows[0].cells[1].text = "分值"
    table.rows[0].cells[2].text = "评分标准"
    table.rows[1].cells[0].text = "检测报告"
    table.rows[1].cells[1].text = "5"
    table.rows[1].cells[2].text = "提供得分"
    document.save(file_path)

    _, parse_result = load_document(file_path)

    assert any(item.node_type.value == "table" for item in parse_result.document_nodes)
    assert any(item.node_type.value == "table_row" for item in parse_result.document_nodes)

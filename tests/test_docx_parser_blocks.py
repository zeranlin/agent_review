from pathlib import Path

from docx import Document

from agent_review.engine import TenderReviewEngine
from agent_review.parsers import load_document


def test_docx_parser_preserves_raw_blocks_and_tables(tmp_path: Path) -> None:
    file_path = tmp_path / "sample.docx"
    document = Document()
    document.add_paragraph("目录", style="Title")
    document.add_paragraph("第一章 招标公告", style="Heading 1")
    document.add_paragraph("一、投标人资格要求", style="Heading 2")
    document.add_paragraph("投标人须具备相关资质。")
    table = document.add_table(rows=2, cols=3)
    table.rows[0].cells[0].text = "评分项"
    table.rows[0].cells[1].text = "分值"
    table.rows[0].cells[2].text = "评分标准"
    table.rows[1].cells[0].text = "检测报告"
    table.rows[1].cells[1].text = "5"
    table.rows[1].cells[2].text = "提供得分"
    document.save(file_path)

    _, parse_result = load_document(file_path)

    assert parse_result.raw_blocks
    assert parse_result.raw_tables
    assert parse_result.raw_blocks[0].metadata["heading_candidate"] is True
    assert parse_result.raw_blocks[0].metadata["catalog_candidate"] is True
    assert any(item.style_name for item in parse_result.raw_blocks)
    assert any(item.numbering for item in parse_result.raw_blocks)
    assert parse_result.raw_tables[0].rows[0][0].is_header is True
    assert parse_result.raw_tables[0].rows[1][2].text == "提供得分"
    assert "第一章 招标公告" in parse_result.text
    assert parse_result.tables


def test_docx_parser_remains_compatible_with_review_engine(tmp_path: Path) -> None:
    file_path = tmp_path / "review.docx"
    document = Document()
    document.add_paragraph("项目概况")
    document.add_paragraph("采购需求：提供家具供货与安装服务。")
    document.add_paragraph("投标人资格要求：具备相关资质。")
    document.add_paragraph("评分标准：综合评分法。")
    document.add_paragraph("付款方式：按合同约定支付。")
    document.save(file_path)

    report = TenderReviewEngine().review_file(file_path)

    assert report.file_info.document_name == "review.docx"
    assert report.parse_result.raw_blocks
    assert isinstance(report.parse_result.raw_tables, list)
